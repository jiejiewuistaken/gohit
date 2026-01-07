from __future__ import annotations

from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from .types import Chunk, Citation, ExtractedDocument


def _chunk_text(text: str, *, max_chars: int = 900) -> list[str]:
    # Simple chunker: split by blank lines then pack.
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf: list[str] = []
    size = 0
    for p in paras:
        if size + len(p) + 2 > max_chars and buf:
            chunks.append("\n\n".join(buf).strip())
            buf = [p]
            size = len(p)
        else:
            buf.append(p)
            size += len(p) + 2
    if buf:
        chunks.append("\n\n".join(buf).strip())
    return [c for c in chunks if c]


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # Basic table extraction (best-effort).
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line.strip():
                parts.append(line)
    return "\n\n".join(parts).strip()


def _read_pdf_pages(path: Path) -> list[str]:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        pages.append(txt)
    return pages


def ingest_file(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = _read_pdf_pages(path)
        chunks: list[Chunk] = []
        for i, page_text in enumerate(pages, start=1):
            for c in _chunk_text(page_text):
                chunks.append(Chunk(text=c, citation=Citation(path=path, page=i)))
        return ExtractedDocument(path=path, kind="pdf", chunks=chunks)

    if suffix == ".docx":
        text = _read_docx(path)
        chunks = [Chunk(text=c, citation=Citation(path=path, page=None)) for c in _chunk_text(text)]
        return ExtractedDocument(path=path, kind="docx", chunks=chunks)

    if suffix in {".txt", ".md"}:
        text = _read_txt(path)
        chunks = [Chunk(text=c, citation=Citation(path=path, page=None)) for c in _chunk_text(text)]
        return ExtractedDocument(path=path, kind=suffix.lstrip("."), chunks=chunks)

    # Fallback: try as text
    text = _read_txt(path)
    chunks = [Chunk(text=c, citation=Citation(path=path, page=None)) for c in _chunk_text(text)]
    return ExtractedDocument(path=path, kind="unknown", chunks=chunks)


def ingest_paths(paths: list[Path]) -> list[ExtractedDocument]:
    docs: list[ExtractedDocument] = []
    for p in paths:
        docs.append(ingest_file(p))
    return docs

