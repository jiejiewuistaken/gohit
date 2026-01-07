from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument

from .types import Chunk


@dataclass(frozen=True)
class CosopMeta:
    country: str
    period_start: str
    period_end: str
    eb_session: str | None = None
    eb_meeting_date: str | None = None
    document_code: str | None = None
    agenda_item: str | None = None
    sec_date: str | None = None
    original_language: str | None = None
    useful_references: str | None = None


def _md_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    out = []
    out.append("| " + " | ".join(header) + " |")
    out.append("| " + " | ".join(["---"] * len(header)) + " |")
    for r in body:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def render_markdown(meta: CosopMeta, *, sections: dict[str, str], evidence_index: list[str]) -> str:
    cover = f"""Executive Board
[{meta.eb_session or 'Insert #'}] Session
Rome, [{meta.eb_meeting_date or 'Insert meeting date'}]

COSOP template-final-e

Document: [{meta.document_code or 'Insert EB../../..'}]
Agenda: [{meta.agenda_item or 'SEC to insert agenda item number'}]
Date: [{meta.sec_date or 'SEC to insert date'}]
Distribution: Public
Original: {meta.original_language or 'Select the original language of the document'}
FOR: REVIEW
Useful references: {meta.useful_references or '[Include useful references separated by commas]'}
Action: The Executive Board is invited to review the country strategic opportunities programme {meta.period_start} to {meta.period_end} for {meta.country}.
"""

    table1 = _md_table(
        [
            ["Indicator", "Data", "Year of reference"],
            ["GNI per capita", "TBD", "TBD"],
            ["GDP growth", "TBD", "TBD"],
            ["Public debt (% of GDP)", "TBD", "TBD"],
            ["Debt service ratio", "TBD", "TBD"],
            ["Debt to GDP ratio", "TBD", "TBD"],
            ["Inflation rate (%)", "TBD", "TBD"],
            ["Population size", "TBD", "TBD"],
            ["Population, female", "TBD", "TBD"],
            ["Youth population", "TBD", "TBD"],
            ["Unemployment rate", "TBD", "TBD"],
            ["Fragility index", "TBD", "TBD"],
            ["INFORM Risk Index", "TBD", "TBD"],
        ]
    )

    table2 = _md_table(
        [
            ["Key development priority (strategic objective)", "Underlying institutions", "Policy reform challenges", "Proposed interventions (lending, non-lending)"],
            ["TBD", "TBD", "TBD", "TBD"],
        ]
    )

    table3 = _md_table(
        [
            ["Project", "Source (PBAS, BRAM)", "IFAD financing", "RTA", "Domestic cofinancing", "International cofinancing", "Cofinancing ratio"],
            ["Ongoing (TBD)", "TBD", "TBD", "TBD", "TBD", "TBD", "TBD"],
            ["Planned (TBD)", "TBD", "TBD", "TBD", "TBD", "TBD", "TBD"],
            ["Total", "", "TBD", "TBD", "TBD", "TBD", "1:XX.X"],
        ]
    )

    rmf = _md_table(
        [
            ["Related UNSDCF/SDG outcomes", "IFAD SOs", "Key COSOP results", "Investments & non-financial activities", "Outcome indicators (targets)", "Output indicators (targets)"],
            ["TBD", "TBD", "TBD", "TBD", "TBD", "TBD"],
        ]
    )

    def s(key: str) -> str:
        return sections.get(key, "").strip() or "_TBD — insufficient evidence in provided inputs._"

    evidence = "\n".join([f"- {e}" for e in evidence_index]) if evidence_index else "- (none)"

    md = f"""{cover}

## Contents
_MVP note: This file is generated and may need formatting for EB submission._

## Executive summary
{s("exec_summary")}

## I. Country context

### A. Socioeconomic setting
{table1}

{s("I.A")}

### B. Transition scenario
{s("I.B")}

### C. Food system, agricultural and rural sector agenda
{s("I.C")}

## II. IFAD engagement: lessons learned

### A. Results achieved during the previous COSOP
{s("II.A")}

### B. Lessons from the previous COSOP and other sources
{s("II.B")}

## III. Strategy for transformational country programmes

### A. COSOP theory of change
{s("III.A")}

### B. Overall goal and strategic objectives
{table2}

{s("III.B")}

### C. Target group and targeting strategy
{s("III.C")}

## IV. IFAD interventions
{s("IV")}

## V. COSOP implementation
### A. Investment volume and sources
{table3}

{s("V")}

## VI. Target group engagement
{s("VI")}

## VII. Risk management
{s("VII")}

## Appendix I. Results management framework (RMF)
{rmf}

## Evidence index (inputs used)
{evidence}
"""
    return md


def render_docx(out_path: Path, markdown_text: str) -> None:
    # MVP: very simple Markdown→DOCX (no full fidelity).
    doc = DocxDocument()
    for line in markdown_text.splitlines():
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("| ") and line.endswith(" |"):
            # Leave tables as preformatted text in MVP.
            doc.add_paragraph(line)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(str(out_path))

